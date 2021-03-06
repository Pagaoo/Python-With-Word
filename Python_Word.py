from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


documento = Document()  # Criando um documento

documento.add_heading('Titulo', 0)  # Adiciona titulo ao documento e o número é o nivel de identação de 0-9

p = documento.add_paragraph('Um paragráfo aleatório ')
p.add_run('Negrito').bold = True  # Adiciona texto no final do paragráfo
p.add_run(' E ')
p.add_run('Italico').italic = True

run = p.add_run('Estilizado')  # Texto estilizado
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alinha o p no centro do word
fonte = run.font
fonte.name = 'Corbel'
fonte.size = Pt(40)
fonte.color.rgb = RGBColor(255, 0, 0)

documento.add_heading('Sub-titulo, nivel 1', level=1)
documento.add_paragraph('Quote', style='Intense Quote')

documento.add_paragraph(
    'Primeiro item na lista', style='List Bullet'
)
documento.add_paragraph(
    'Primeiro item na lista numerada', style='List Number'
)

documento.add_paragraph(
    'Segundo item na lista numerada', style='List Number'
)

documento.add_picture('Xiao.jpg', width=Inches(2))

# Criação de uma tupla com items da tabela
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

# Cria uma tabela com as especificações a baixo
table = documento.add_table(rows=1, cols=3)  # Define quantas colunas e quantas filas criar
hdr_cells = table.rows[0].cells  # Cria a tabela a partir do index 0
hdr_cells[0].text = 'Qty'  # Cria uma célula com um texto de 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:  # Para cada item adiciona uma celula
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

documento.add_page_break()  # Dá um break na página

documento.save('arquivo.docx')  # Salva o documento
